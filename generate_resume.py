"""
Generate an ATS-friendly Word document resume for Rohan Kadam.
Tailored for Williams F1 Systems Simulation Engineer role.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

# ============================================================
# HEADER — Name & Contact
# ============================================================
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name.add_run('ROHAN KADAM')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
name.paragraph_format.space_after = Pt(2)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Systems Simulation Engineer — Powertrain, Hydraulics & 1D Modeling')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
title.paragraph_format.space_after = Pt(6)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('+91 96734 75039  |  rohankadam.1023@gmail.com  |  Pune, India  |  linkedin.com/in/rohan-kadam-7bb932161  |  github.com/rohaaaaaan')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
contact.paragraph_format.space_after = Pt(4)


def add_section_heading(text):
    """ATS-friendly section heading with a bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '333333',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ============================================================
# PROFESSIONAL SUMMARY
# ============================================================
add_section_heading('Professional Summary')
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(6)
run = summary.add_run(
    'Systems Simulation Engineer with over 6 years of experience specializing in 1D physics-based simulation, '
    'model-to-track correlation, and physical systems engineering (hydraulics, cooling, powertrains, and braking). '
    'Proven track record of using simulation outputs (transient thermal, mechanical matching, and kinematics) to '
    'directly guide design concepts, select suppliers, and shape vehicle performance targets. Highly proficient in '
    'MATLAB/Simulink and Python for telemetry data analysis, with strong familiarity in Modelica-based 1D physical '
    'modeling (Dymola) and thermal simulators (GT-Suite). Strong first-principles engineering background rooted in '
    'Formula Student motorsport developments and production-scale vehicle validation.'
)
run.font.size = Pt(10)


# ============================================================
# CORE COMPETENCIES (keyword-rich for ATS)
# ============================================================
add_section_heading('Core Competencies')
competencies = doc.add_paragraph()
competencies.paragraph_format.space_after = Pt(6)
run = competencies.add_run(
    '1D Physics-Based Simulation (MATLAB/Simulink, Dymola/Modelica, GT-Suite)  •  Model Correlation & Track/Test Data Validation  •  '
    'Mechanical & Fluid Systems Sizing (Hydraulics, cooling, powertrain, braking)  •  Python Data Analysis & Telemetry Log Processing  •  '
    'First-Principles Engineering Design & Pugh Decision Analysis  •  Transient Thermal & Fluid Dynamics (CFD, Heat Transfer)  •  '
    'Hardware-in-the-Loop (HIL) Testing & Sensor Integration  •  CAD Packaging & PLM (CATIA V5, Siemens Teamcenter)'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ============================================================
# PROFESSIONAL EXPERIENCE
# ============================================================
add_section_heading('Professional Experience')

experiences = [
    {
        'company': 'Cognizant (Client: Stellantis)',
        'title': 'Systems Simulation & Controls Engineer — EV Range & Powerflow',
        'period': '2024 – Present',
        'location': 'Pune, India',
        'bullets': [
            'Developed and validated dynamic 1D physics-based energy models and control algorithms in MATLAB/Simulink to simulate EV range estimation (Distance-to-Empty) and real-time Powerflow visualization.',
            'Built a custom Python-based telemetry processing engine (Advanced Signal Data Analysis Tool) to automate post-processing and alignment of HIL and track road test logs, identifying anomalies and correlating simulations to real-world datasets.',
            'Managed feature specifications and E/E architecture requirements (DOORS Next Gen, ASPICE) for EV range algorithms, ensuring system behavior matches vehicle-level targets across next-gen STLA platforms.',
            'Deployed Model-Based Systems Engineering (MBSE) using SysML (Rhapsody, Cameo/Magic) to design functional interfaces and port contracts between physical powertrain, thermal loops, and control modules.',
            'Executed Hardware-in-the-Loop (HIL) calibration and checkout tests, utilizing Vector CANoe and ETAS INCA to debug interface parameters and analyze dynamic bus traffic (CAN FD/Ethernet).',
            'Owned ISO 26262 functional safety requirements up to ASIL-D for torque arbitration and powertrain securement paths, designing safe-state torque-neutral transitions.',
        ]
    },
    {
        'company': 'Mahindra & Mahindra',
        'title': 'Thermal & Fluid Systems Simulation Engineer',
        'period': '2023 – 2024',
        'location': 'Pune, India',
        'bullets': [
            'Primary owner of vehicle cooling and active thermal management loops; utilized 1D GT-Suite simulation models to size radiator and intercooler cores based on Logarithmic Mean Temperature Difference (LMTD) and convective heat transfer equations.',
            'Led dynamic wind tunnel validation and field testing; set up high-density thermocouple grids and hot-wire anemometers on prototype vehicles, correlating experimental results with 1D/3D models to refine boundary conditions.',
            'Translated simulation findings into engineering decisions: analyzed underhood 3D CFD velocity and temperature profiles to eliminate recirculation zones, designing custom airflow deflectors and heat shields.',
            'Defined requirements and control logic for a 5-way rotary electric coolant valve to dynamically configure battery/e-motor cooling loops (serial, parallel, and waste-heat recovery modes).',
            'Developed 3D CAD packaging and routing layouts for high-pressure fluid lines using CATIA V5, releasing production drawings via Siemens Teamcenter PLM.',
            'Led design reviews and FMEA sessions to secure the ISO 26262 ASIL-C safety case for the accelerator pedal module, resolving 15+ packaging bottlenecks as VP0 Build Champion.',
        ]
    },
    {
        'company': 'TATA Technologies',
        'title': 'Powertrain Simulation & Design Engineer — Forklift Systems',
        'period': '2019 – 2023',
        'location': 'Pune, India',
        'bullets': [
            'Developed a first-principles 1D simulation model in Excel and Python for engine-to-torque converter matching, matching engine curves with converter K-factors and ratio multipliers to predict tractive effort and gradeability (mean error within 2% of track data).',
            'Employed weighted Pugh Decision Matrices to connect torque matching simulation outputs to engineering decisions, guiding selection of engine-transmission packages for 70+ forklift variants.',
            'Conducted 1D thermal and fluid flow simulations for radiator/oil cooler packages and sized mechanical/electric cooling fans to optimize heat dissipation against parasitic engine load.',
            'Configured J1939 CAN bus network message routing for fleet telematics and configured local diagnostic trouble code (DTC) alert logic for system-level faults.',
            'Collaborated with tier-1 engine suppliers to redesign flywheel couplings, resolving critical input shaft packaging mismatches discovered during prototype builds.',
        ]
    },
]

for exp in experiences:
    # Company + Period on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(exp['company'])
    run.bold = True
    run.font.size = Pt(11)
    if exp['period']:
        run = p.add_run(f"  —  {exp['period']}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Title + Location
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(exp['title'])
    run.italic = True
    run.font.size = Pt(10.5)
    run = p.add_run(f"  |  {exp['location']}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for bullet in exp['bullets']:
        add_bullet(bullet)


# ============================================================
# EDUCATION / EARLY EXPERIENCE (Motorsport Focus)
# ============================================================
add_section_heading('Education & Motorsport Experience')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Veloce Racing — Formula Student Team')
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('Brake Systems Design & Simulation Lead')
run.italic = True
run.font.size = Pt(10.5)

add_bullet('Designed hydraulic brake control circuits and sized master cylinder/caliper bores from first principles, balancing weight transfer and deceleration physics (MATLAB/Simulink).')
add_bullet('Designed and manufactured an adjustable mechanical brake balance bar, allowing real-time front-to-rear pressure bias tuning based on track test feedback.')
add_bullet('Performed transient thermal stress finite element analysis (FEA) on ventilated brake rotors in ANSYS; validated convective heat transfer assumptions against track telemetry via surface thermocouple sensors.')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Bachelor of Engineering — Mechanical Engineering')
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('Vishwakarma Institute of Technology (VIT), Savitribai Phule Pune University  |  2019  |  CGPA: 8.02')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('HSC (12th Standard)')
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run('  |  2015  |  84.2%')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('SSC (10th Standard)')
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run('  |  2013  |  93.46%')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ============================================================
# KEY PROJECTS
# ============================================================
add_section_heading('Key Projects')

projects = [
    {
        'name': 'AI-Powered Systems Engineering Platform (AISE)',
        'tech': 'Python, LangChain, React, MBSE, SysML, RAG',
        'desc': 'Intelligent systems engineering platform automating requirements decomposition and trace matrices using LLMs. Converts natural language requirements into SysML-compliant block definitions and interface contracts.',
    },
    {
        'name': 'F1 Telemetry & 3D Race Visualisation',
        'tech': 'React, Three.js, Node.js, WebSockets, MQTT, Python',
        'desc': 'A real-time telemetry streaming and interactive 3D visualization dashboard. Decodes live CAN and sensor streams over WebSockets to render vehicle kinetics, tire thermal zones, and track positioning in a high-performance WebGL-based 3D environment.',
    },
]

for proj in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(proj['name'])
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(f"  [{proj['tech']}]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph(proj['desc'])
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)


# ============================================================
# TECHNICAL SKILLS (categorized for ATS)
# ============================================================
add_section_heading('Technical Skills')

skills = [
    ('Simulation & 1D Modeling', 'MATLAB / Simulink, Dymola / Modelica (Familiar), GT-Suite (1D Thermal), ANSYS (Thermal FEA), Simscape'),
    ('Languages & Analysis', 'Python (Pandas, NumPy, Matplotlib), C++, SQL, Git, Telemetry Post-Processing'),
    ('Testing & Telemetry', 'Vector CANoe, ETAS INCA, dSPACE HIL, Wind Tunnel Testing, Thermocouple & Hot-wire Anemometer Rigs'),
    ('CAD & PLM & GD&T', 'CATIA V5 (Packaging & Routing), SolidWorks, Siemens Teamcenter PLM, GD&T'),
    ('Systems & Architecture', 'MBSE / SysML (IBM Rhapsody, Cameo/Magic), AUTOSAR Classic & Adaptive, ISO 26262 (ASIL-C/D), ASPICE, APQP / FMEA, 8D Problem Solving'),
    ('Protocols & Networks', 'CAN / CAN FD, LIN, FlexRay, Automotive Ethernet, J1939'),
]

for category, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{category}: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(items)
    run.font.size = Pt(10)


# ============================================================
# SAVE & CONVERT TO PDF
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'public', 'Rohan_Kadam_Resume_v2.docx')
doc.save(output_path)
print(f'Resume saved to: {output_path}')

# Compile PDF version for the website
try:
    import win32com.client
    print("Converting to PDF...")
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    pdf_path = os.path.join(os.path.dirname(__file__), 'public', 'Rohan Kadam Systems Engineer.pdf')
    docx_doc = word.Documents.Open(os.path.abspath(output_path))
    docx_doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF
    docx_doc.Close()
    word.Quit()
    print(f'PDF compiled successfully to: {pdf_path}')
except Exception as e:
    print(f'Warning: Could not automatically compile PDF. Error: {e}')

